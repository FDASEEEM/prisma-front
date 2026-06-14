"""Conversor mínimo Markdown -> .docx para el INFORME de P.R.I.S.M.A.

Soporta: encabezados (#..####), párrafos, listas (- / 1.), tablas GFM,
bloques de código (```), citas (>), negritas **..** e inline `code`.
No pretende ser un parser completo; cubre lo que usa INFORME.md.
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC, OUT = sys.argv[1], sys.argv[2]

with open(SRC, encoding="utf-8") as fh:
    lines = fh.read().splitlines()

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]+)`")


def add_runs(paragraph, text):
    """Renderiza **negrita** e `inline code` dentro de un párrafo."""
    # Partimos primero por negritas, luego por inline-code dentro de cada trozo.
    pos = 0
    tokens = []
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            tokens.append(("plain", text[pos:m.start()]))
        tokens.append(("bold", m.group(1)))
        pos = m.end()
    if pos < len(text):
        tokens.append(("plain", text[pos:]))

    for kind, chunk in tokens:
        sub = 0
        for cm in CODE_RE.finditer(chunk):
            if cm.start() > sub:
                _run(paragraph, chunk[sub:cm.start()], kind == "bold", False)
            _run(paragraph, cm.group(1), kind == "bold", True)
            sub = cm.end()
        if sub < len(chunk):
            _run(paragraph, chunk[sub:], kind == "bold", False)


def _run(paragraph, text, bold, code):
    r = paragraph.add_run(text)
    r.bold = bold
    if code:
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)


def add_table(rows):
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [
        [c.strip() for c in r.strip("|").split("|")]
        for r in rows[2:]  # rows[1] es el separador |---|
    ]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        add_runs(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if i < len(cells):
                cells[i].paragraphs[0].text = ""
                add_runs(cells[i].paragraphs[0], val)


i = 0
while i < len(lines):
    line = lines[i]

    # Bloque de código ```
    if line.strip().startswith("```"):
        i += 1
        buf = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            buf.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        r = p.add_run("\n".join(buf))
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        i += 1
        continue

    # Tabla GFM (línea con | y la siguiente con |---)
    if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").strip()) <= set("-: "):
        block = []
        while i < len(lines) and lines[i].startswith("|"):
            block.append(lines[i])
            i += 1
        add_table(block)
        continue

    stripped = line.strip()

    if not stripped:
        i += 1
        continue
    if stripped.startswith("####"):
        doc.add_heading(stripped.lstrip("#").strip(), level=4)
    elif stripped.startswith("###"):
        doc.add_heading(stripped.lstrip("#").strip(), level=3)
    elif stripped.startswith("##"):
        doc.add_heading(stripped.lstrip("#").strip(), level=2)
    elif stripped.startswith("#"):
        h = doc.add_heading(stripped.lstrip("#").strip(), level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif stripped.startswith("---"):
        pass  # separador horizontal -> lo omitimos
    elif stripped.startswith(">"):
        p = doc.add_paragraph(style="Intense Quote")
        add_runs(p, stripped.lstrip(">").strip())
    elif re.match(r"^\d+\.\s", stripped):
        p = doc.add_paragraph(style="List Number")
        add_runs(p, re.sub(r"^\d+\.\s", "", stripped))
    elif stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, stripped[2:])
    else:
        p = doc.add_paragraph()
        add_runs(p, stripped)
    i += 1

doc.save(OUT)
print(f"OK -> {OUT}")
